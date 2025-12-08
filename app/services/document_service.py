from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generar_pdf(documento_texto: str, nombre_solicitante: str = "Documento") -> BytesIO:
    """
    Genera un PDF del documento legal
    """
    buffer = BytesIO()

    # Crear el documento PDF
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )

    # Estilos
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='Justify',
        parent=styles['BodyText'],
        alignment=TA_JUSTIFY,
        fontSize=11,
        leading=14,
    ))

    # Contenido
    story = []

    # Dividir el texto en párrafos
    lineas = documento_texto.split('\n')

    for linea in lineas:
        if linea.strip():
            # Detectar títulos (líneas que empiezan con ** o están en mayúsculas)
            if linea.strip().startswith('**') or linea.strip().isupper():
                # Es un título
                texto_limpio = linea.replace('**', '').strip()
                p = Paragraph(f"<b>{texto_limpio}</b>", styles['Heading2'])
            else:
                # Es texto normal
                p = Paragraph(linea, styles['Justify'])

            story.append(p)
            story.append(Spacer(1, 0.1 * inch))
        else:
            # Línea vacía - agregar espacio
            story.append(Spacer(1, 0.2 * inch))

    # Construir el PDF
    doc.build(story)
    buffer.seek(0)

    return buffer


def generar_docx(documento_texto: str, nombre_solicitante: str = "Documento") -> BytesIO:
    """
    Genera un archivo DOCX del documento legal
    """
    buffer = BytesIO()

    # Crear documento
    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Dividir el texto en líneas
    lineas = documento_texto.split('\n')

    for linea in lineas:
        if linea.strip():
            # Detectar títulos
            if linea.strip().startswith('**') or (len(linea.strip()) > 0 and linea.strip().isupper() and len(linea.strip()) < 100):
                # Es un título
                texto_limpio = linea.replace('**', '').strip()
                heading = doc.add_heading(texto_limpio, level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # Es texto normal
                p = doc.add_paragraph(linea)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Configurar fuente
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
        else:
            # Línea vacía
            doc.add_paragraph()

    # Guardar en buffer
    doc.save(buffer)
    buffer.seek(0)

    return buffer
